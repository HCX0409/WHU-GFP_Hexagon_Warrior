"""
生成项目详细文档 (DOC格式) - v4 (基于代码实现)
包含文本流程图、目录、AlphaFold3分析、模型来源汇总
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ======================== 全局样式 ========================
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = style.paragraph_format
pf.space_after = Pt(4)
pf.space_before = Pt(2)

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    if i == 1:
        hs.font.size = Pt(16)
    elif i == 2:
        hs.font.size = Pt(14)
    else:
        hs.font.size = Pt(12)

def add_para(text, bold=False, indent=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.75)
    return p

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_table(headers, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri+1].cells[ci].text = str(val)
    return t

def add_flowbox(text, is_last=False):
    """用带边框的段落模拟流程图节点"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'  {text}  ')
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(11)
    run.bold = True
    # 加边框
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['top', 'left', 'bottom', 'right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '8')
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), '2E74B5')
        pBdr.append(el)
    pPr.append(pBdr)
    # 箭头
    if not is_last:
        arrow = doc.add_paragraph()
        arrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar = arrow.add_run('|\nV')
        ar.font.size = Pt(12)
        ar.bold = True
        ar.font.color.rgb = RGBColor(46, 116, 181)

# ======================== 封面 ========================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\nGFP亮度优化设计项目\n详细技术文档')
run.font.size = Pt(22)
run.bold = True
run.font.name = 'Times New Roman'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p2.add_run('\n基于ESM-2蛋白质语言模型的LoRA微调与多模型协同策略\n\nGFP Hexagon Warrior — SynBio Challenges 2026\nWHU-CHINA\n作者: HCX')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ======================== 目录 ========================
add_heading('目录', level=1)
toc_items = [
    ('1.', '项目概述'),
    ('2.', '项目整体流程'),
    ('  2.1', '项目思维导图'),
    ('3.', '数据基础与预处理'),
    ('  3.1', '原始数据'),
    ('  3.2', '数据清洗'),
    ('  3.3', '序列解析与Delta策略'),
    ('  3.4', '72k优质数据筛选'),
    ('4.', '模型架构演进 (V1-V9)'),
    ('  4.1', '早期探索: V1-V4 (传统ML)'),
    ('  4.2', 'LoRA微调: V5'),
    ('  4.3', '高亮度聚焦: V6'),
    ('  4.4', '冻结主干: V7 / V7.5'),
    ('  4.5', '实验尝试: V8 Gated MLP'),
    ('  4.6', '最终主力: V9'),
    ('  4.7', '三模型协同体系'),
    ('5.', 'Phase 2: 分层突变与迭代筛选'),
    ('  5.1', '父本选取与突变生成'),
    ('  5.1b', '早期探索性突变扫描'),
    ('  5.2', '两阶段筛选 (V7初筛 + V9/V5精筛)'),
    ('  5.3', '多轮迭代与累计合并'),
    ('6.', 'Phase 3: 热稳定性分析 (ESM3-ddG)'),
    ('7.', 'Phase 4: 最终候选筛选'),
    ('8.', 'Phase 5: AlphaFold 3 结构审核'),
    ('9.', '最终候选序列'),
    ('10.', '终极汇总管线 (final_pipeline)'),
    ('11.', '模型与工具来源'),
    ('12.', '辅助工具与验证程序'),
    ('13.', '项目文件清单'),
    ('14.', '参考文献'),
]
for num, title in toc_items:
    add_para(f'{num} {title}')

doc.add_page_break()

# ======================== 1. 项目概述 ========================
add_heading('1. 项目概述', level=1)
add_para('本项目参加SynBio Challenges 2026合成生物学竞赛,目标是设计6条高亮度绿色荧光蛋白(GFP)变体序列。核心方法:利用蛋白质语言大模型ESM-2 (650M)的LoRA微调技术构建三模型协同预测系统,结合ESM3-ddG热稳定性验证与AlphaFold 3人工结构审核,从12万条实验数据中挖掘高亮度突变模式。', indent=True)

# ======================== 2. 项目流程图 ========================
add_heading('2. 项目整体流程', level=1)
add_para('以下为项目从原始数据到最终候选序列的完整五阶段流程:', indent=True)

# 插入流程图
doc.add_picture('d:/GFP_Hexagon_Warrior/assets/流程图.png', width=Cm(15))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption = doc.add_paragraph()
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = caption.add_run('图1: 项目整体五阶段流程图')
run.font.size = Pt(10)
run.italic = True
run.font.name = 'Times New Roman'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 插入思维导图
add_heading('2.1 项目思维导图', level=2)
doc.add_picture('d:/GFP_Hexagon_Warrior/assets/项目导图.png', width=Cm(15))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption = doc.add_paragraph()
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = caption.add_run('图2: 项目思维导图')
run.font.size = Pt(10)
run.italic = True
run.font.name = 'Times New Roman'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ======================== 3. 数据基础 ========================
add_heading('3. 数据基础与预处理', level=1)
add_para('对应程序: src/phase1_data.py, src/prepare_72k_data.py', indent=False)

add_heading('3.1 原始数据', level=2)
add_para('GFP_data.xlsx中brightness表,141,572条实验记录,4种GFP类型:', indent=True)
add_table(
    ['GFP类型', 'WT亮度(log10)', '数据量', '超高亮度(>4.3)'],
    [
        ['avGFP', '3.72', '51,715', '0%'],
        ['amacGFP', '3.97', '33,809', '0%'],
        ['cgreGFP', '4.50', '24,568', '41.7%'],
        ['ppluGFP', '4.23', '31,480', '4.5%'],
    ]
)

add_heading('3.2 数据清洗', level=2)
add_para('(1) 删除亮度完全相同的异常数据(出现次数超过50次的死值,不符合实验测量误差规律)及负数亮度。', indent=True)
add_para('(2) 过滤含终止密码子(*)的无效突变及解析失败序列。')
add_para('(3) 序列去重(不同突变字符串可能对应相同氨基酸序列),最终有效训练数据122,836条。')

add_heading('3.3 序列解析', level=2)
add_para('原始数据以突变标注格式存储(如"A109D:K130M"),需解析为完整238位氨基酸序列。亮度预测采用delta策略: 预测目标 = 实验亮度 - WT亮度,消除不同GFP类型间的系统偏差。', indent=True)

add_heading('3.4 72k优质数据筛选', level=2)
add_para('全量122,836条数据分布两极分化(大量死蛋白+少量高亮度)。V7/V9采用分层采样筛选72,000条(58.6%),亮度从低到高均匀覆盖,按GFP类型独立处理。', indent=True)

# ======================== 4. 模型架构演进 ========================
add_heading('4. 模型架构演进 (V1-V9)', level=1)
add_para('对应程序: src/phase1_lora_finetune_v5.py (V5), src/phase1_frozen_esm2_mlp.py (V7), src/phase1_lora_finetune_v6.py (V9)', indent=False)
add_para('项目经历9个版本迭代,从传统机器学习到深度学习微调:', indent=True)

add_heading('4.1 早期探索: V1-V4 (传统机器学习)', level=2)
add_table(
    ['版本', '方法', 'R2', '淘汰原因'],
    [
        ['V1', 'ESM-2嵌入 + XGBoost', '~0.55', '基线模型,精度有限'],
        ['V2-V3', '增加物理化学特征', '~0.58', '提升有限'],
        ['V4', 'XGBoost参数优化', '~0.62', '高亮度区间偏差大'],
    ]
)
add_para('核心问题: 冻结嵌入无法通过端到端训练适应亮度预测任务。', indent=True)

# V4预测散点图
add_para('V4模型各GFP类型预测散点图:', indent=False)
for gfp in ['avGFP', 'amacGFP', 'cgreGFP', 'ppluGFP']:
    img_path = f'd:/GFP_Hexagon_Warrior/prediction comparison/model_generalization_v4_{gfp}.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Cm(13))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# V4散点图R²标注
caption = doc.add_paragraph()
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = caption.add_run('V4分类型R²: avGFP=0.840, amacGFP=0.555, cgreGFP=0.476, ppluGFP=0.461 (各N=1000均匀采样)')
run.font.size = Pt(9)
run.italic = True
run.font.name = 'Times New Roman'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

add_heading('4.2 LoRA微调: V5', level=2)
add_para('首次引入LoRA微调,全量122,836条数据训练:', indent=True)
add_para('  架构: ESM-2 t33 650M + LoRA(r=32, alpha=64) + Mean Pooling + GFP类型嵌入(4维) + 回归头(1284->768->512->256->1)')
add_para('  性能: 分类型R²: avGFP=0.816, amacGFP=0.549, cgreGFP=0.442, ppluGFP=0.663。avGFP预测能力最强,cgreGFP最弱。')
add_para('  定位: 三模型体系中权重15%,作为"安全网"覆盖全亮度范围。')

# V5预测散点图
add_para('V5模型预测散点图(各GFP类型):', indent=False)
for gfp in ['avGFP', 'amacGFP', 'cgreGFP', 'ppluGFP']:
    img_path = f'd:/GFP_Hexagon_Warrior/prediction comparison/model_generalization_v5_{gfp}.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Cm(13))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# V5散点图R²标注
caption = doc.add_paragraph()
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = caption.add_run('V5分类型R²: avGFP=0.816, amacGFP=0.549, cgreGFP=0.442, ppluGFP=0.663 (各N=1000均匀采样)')
run.font.size = Pt(9)
run.italic = True
run.font.name = 'Times New Roman'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

add_heading('4.3 高亮度聚焦: V6', level=2)
add_para('构建36,405条高亮度聚焦训练集(>4.0占55.1%),作为"高亮度专家"。后续被V9体系取代,但其聚焦策略被72k数据筛选继承。', indent=True)

add_heading('4.4 冻结主干: V7', level=2)
add_para('冻结ESM-2全部参数,仅训练MLP头(1284->512->256->128->1):', indent=True)
add_para('  优势: 冻结ESM-2,仅训练MLP头,体量小,推理速度快。')
add_para('  性能: 全局R²=0.921 (72k全量数据),分类型R²: avGFP=0.864, amacGFP=0.851, cgreGFP=0.807, ppluGFP=0.880。三模型体系中权重35%,用作初筛。')

# V7预测散点图
add_para('V7模型预测散点图:', indent=False)
if os.path.exists('d:/GFP_Hexagon_Warrior/assets/V7预测能力.png'):
    doc.add_picture('d:/GFP_Hexagon_Warrior/assets/V7预测能力.png', width=Cm(14))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run('V7全局R²=0.921, r=0.960, MAE=0.161 (72k全量数据均匀采样 N=2664)')
    run.font.size = Pt(9)
    run.italic = True
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

add_heading('4.5 实验性尝试: V8与V7.5', level=2)
add_table(
    ['版本', '方法', '结果'],
    [
        ['V8', 'Gated MLP(门控机制)', '过拟合,未优于V7,淘汰'],
        ['V7.5', 'Residual MLP(残差连接)', '与V7接近,更简洁但收益不大'],
    ]
)

add_heading('4.6 最终主力: V9', level=2)
add_para('V5架构 + 72k优质数据 = 精度最高的主力模型:', indent=True)
add_para('  性能: 全局R²=0.946 (72k全量数据),分类型R²: avGFP=0.917, amacGFP=0.891, cgreGFP=0.832, ppluGFP=0.891。各类型预测一致性最好。')
add_para('  定位: 三模型体系中权重50%,排名核心。')

# V9预测散点图
add_para('V9模型预测散点图:', indent=False)
if os.path.exists('d:/GFP_Hexagon_Warrior/assets/v9_scatter_uniform.png'):
    doc.add_picture('d:/GFP_Hexagon_Warrior/assets/v9_scatter_uniform.png', width=Cm(14))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run('V9全局R²=0.946, r=0.973, MAE=0.148 (72k全量数据均匀采样 N=2664)')
    run.font.size = Pt(9)
    run.italic = True
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

add_heading('4.7 三模型体系总结', level=2)
add_table(
    ['模型', '架构', '训练数据', '权重', 'R2', '角色'],
    [
        ['V9', 'ESM-2 LoRA', '72k优质', '50%', '0.946', '精度最高,排名核心'],
        ['V7', 'Frozen MLP', '72k优质', '35%', '0.921', '体量小,推理速度快,用作初筛'],
        ['V5', 'ESM-2 LoRA', '全量122k', '15%', '-', '覆盖最广,识别死蛋白'],
    ]
)
add_para('融合公式: weighted_score = 0.50 x V9 + 0.35 x V7 + 0.15 x V5', indent=True)

# ======================== 5. Phase 2: 分层突变与迭代筛选 ========================
add_heading('5. Phase 2: 分层突变与迭代筛选', level=1)
add_para('对应程序: src/phase2_generate_hierarchical.py', indent=False)
add_para('Phase 2是项目的核心序列生成阶段,采用分层突变策略从已知高亮度序列出发,通过大规模突变生成与多轮迭代筛选,挖掘更高亮度的突变变体。', indent=True)

add_heading('5.1 父本选取与突变生成', level=2)
add_para('从GFP_data.xlsx的brightness表中,按4种GFP类型(avGFP/amacGFP/cgreGFP/ppluGFP)分别选取训练数据中亮度最高的50条序列作为父本,4类型 × 50条 = 200条父本。', indent=True)
add_para('对每条父本序列,在保护区(发色团TYG 65-67位 + 15个核心结构位点)之外的位点进行随机氨基酸替换,每条父本生成500个突变变体,每个变体固定3个突变位点。5轮迭代(每轮不同随机种子)共计生成500,000条原始突变序列(200父本 × 500变体 × 5轮)。')

add_heading('5.1b 早期探索性突变扫描', level=2)
add_para('在分层突变方案确定之前,项目还执行了一轮探索性突变扫描:从各类型Top 1000高亮度序列出发,每条生成20个变体(突变数随机1-5个),共产生80,000条原始变体。使用V5 LoRA模型打分,过滤掉亮度下降超过0.3的死蛋白,去重后保留32,394条(含父本对照)。该扫描的目的是广泛探索突变空间,为后续分层方案提供参考。这些序列后续也纳入了最终汇总管线统一打分。', indent=True)
add_para('对应程序: archive/src_legacy/phase2_generate.py (已归档)', indent=False)

add_heading('5.2 两阶段筛选', level=2)
add_para('生成的500,000条原始序列需经过严格的两阶段筛选:', indent=True)

add_table(
    ['筛选阶段', '方法', '输入', '输出', '通过率'],
    [
        ['V7初筛', 'Frozen MLP快速预测', '500,000', '8,000/轮', '1.6%'],
        ['V9+V5精筛', 'LoRA精确打分+共识投票', '8,000/轮', '数千条', '10-30%'],
        ['5轮合并去重', '累计合并+序列去重', '全部迭代产出', '6,626', '-'],
    ]
)

add_para('  V7初筛: 使用V7 Frozen MLP模型对所有变体进行快速亮度预测,各GFP类型设定独立阈值(cgreGFP>=4.40, ppluGFP>=4.13, amacGFP>=3.80, avGFP>=3.50),淘汰明显低于阈值的变体。每轮500,000条中仅1-2%通过。')
add_para('  V9+V5精筛: 对V7通过的变体,依次用V9 LoRA(精度最高,全局R²=0.946)和V5 LoRA(覆盖最广)精确打分。采用三模型共识机制: 至少2/3的模型预测达到该类型共识目标(如cgreGFP>=4.55),才视为通过。')

add_heading('5.3 多轮迭代与累计合并', level=2)
add_para('整个Phase 2执行5轮迭代,每轮使用不同的随机种子生成全新的变体集合(避免重复探索同一突变空间)。各轮迭代独立筛选后,将通过V7初筛的序列累计合并、去重,形成最终候选序列池。', indent=True)

add_table(
    ['迭代', '新增V7通过', '累计候选池', '主要GFP类型覆盖'],
    [
        ['第1轮', '500', '500', 'ppluGFP, cgreGFP'],
        ['第2轮', '500', '1,000', 'ppluGFP, cgreGFP'],
        ['第3轮', '500', '1,500', 'ppluGFP, cgreGFP'],
        ['第4轮', '2,500', '4,000', 'amacGFP, avGFP新增'],
        ['第5轮', '2,626', '6,626', '全类型覆盖'],
    ]
)
add_para('经过5轮完整迭代,从500,000条原始突变序列中筛选出6,626条通过三模型共识的高质量候选序列,整体筛选率1.3%。', indent=True)

# ======================== 6. 热稳定性分析 ========================
add_heading('6. Phase 3: 热稳定性分析 (ESM3-ddG)', level=1)
add_para('对应程序: src/phase3_ddg_scoring.py', indent=False)
add_para('使用ESM3-ddG v2预测突变自由能变化(ddG):', indent=True)
add_para('  ddG < 0: 稳定化突变(有利结构)')
add_para('  ddG > 0: 去稳定化突变')
add_para('技术限制: 仅支持单突变输入,多突变需逐位点预测后求和。')
add_para('亮度-稳定性关系: Pearson r=0.175(弱正相关)。死蛋白反而最稳定(结构完好但不发光),超亮序列稳定性略低。因此ddG仅作为辅助过滤指标。', indent=True)

# ======================== 7. AlphaFold 3 人工审核 ========================
add_heading('7. Phase 4: 最终候选筛选', level=1)
add_para('对应程序: src/phase4_scoring.py', indent=False)
add_para('Phase 4对Phase 3打分后的候选序列进行最终筛选:', indent=True)
add_para('(1) Exclusion List过滤: 排除135,414条已知序列(比赛规则要求不得与已有序列重复)。')
add_para('(2) 综合评分: 采用z-score加权, 0.3 x 亮度z-score + 0.7 x 稳定性z-score(-ddG),稳定性权重更高。')
add_para('(3) 多样性配额: 每种GFP类型至少保留3条,剩余名额按综合分从高到低填满,总计选出Top 20。')

# ======================== 8. AlphaFold 3 ========================
add_heading('8. Phase 5: AlphaFold 3 人工结构审核', level=1)

add_heading('8.1 审核目的', level=2)
add_para('AI模型预测的亮度值不能盲目信任(R2不等于真实准确率)。最终候选必须经过人工在AlphaFold 3上审核结构完整性,确保:', indent=True)
add_para('(1) beta-barrel桶状结构完整,折叠正确。')
add_para('(2) 发色团(TYG)区域稳定(pLDDT不低于70)。')
add_para('(3) 在72度热处理10分钟+降温至25度的比赛条件下,蛋白质不会因结构不稳定而变性或聚集。')

add_heading('8.2 先期研究: 亮度-结构关系', level=2)
add_para('为建立审核标准,我们对GFP_data.xlsx中各类型GFP按亮度分高/中/低三档,各取10条(共120条)进行AlphaFold 3结构预测,核心发现:', indent=True)
add_para('(1) 高亮度序列: 发色团区域pLDDT通常在70以上(深蓝至浅蓝),核心结构稳定。')
add_para('(2) 低亮度序列: 发色团区域pLDDT常低于70(出现黄色/橙色),结构不稳定。')
add_para('(3) 重要例外: 部分序列AlphaFold 3预测"完美"(全结构pLDDT>70),但实际亮度很低。这说明结构稳定是发光的必要条件,但不充分——亮度还取决于发色团微环境的电子结构,这正好是AI模型需要预测的部分。')

add_heading('8.3 审核流程', level=2)
add_para('最终候选的审核流程:', indent=True)
add_para('(1) 三模型AI交叉验证: 三个模型均认为亮 -> 真实亮概率大。')
add_para('(2) AlphaFold 3结构预测: 检查beta-barrel完整性、发色团pLDDT。')
add_para('(3) 参考标准: 以mBaoJin GFP(目前已知最亮单体荧光蛋白)作为结构参考。')
add_para('(4) 热稳定性判断: 确保在72度/10min+25度条件下结构不崩塌。')

add_heading('8.4 审核结果与替换', level=2)
add_para('v1初版中2条候选未通过审核:', indent=True)
add_para('(1) avGFP候选(N146I, K162A): beta-barrel区域结构不稳定,替换。')
add_para('(2) cgreGFP候选: 144-148区域beta折叠不完整,桶状结构有问题,替换。')
add_para('经v2、v3两轮替换迭代,从Top 10候选列表中逐一尝试,最终v3版本全部5条通过结构审核。')

# ======================== 9. 最终候选 ========================
add_heading('9. 最终候选序列', level=1)
add_para('将4个数据源的候选序列合并,去重后39,020条统一三模型打分:', indent=True)
add_table(
    ['数据源', '来源', '序列数', '突变数分布'],
    [
        ['分层突变(Phase 2)', '5轮迭代筛选', '6,626', '固定3个'],
        ['探索性扫描', '早期随机突变', '32,394', '0-5个(随机)'],
        ['交叉验证', '训练数据Top序列', '1,597', '0-5个'],
        ['Phase 4 Top 20', 'ddG筛选后', '20', '3-5个'],
    ]
)
add_para('合并去重后39,020条统一三模型打分。选取规则: 4种GFP类型各至少1条 + 第5条给最高分类型,通过Exclusion List过滤。经AlphaFold 3结构审核(v1~v3迭代),最终确认5条候选序列。', indent=True)

add_table(
    ['序号', 'GFP类型', '加权亮度', '关键突变', 'ddG', '来源'],
    [
        ['1', 'cgreGFP', '4.464', 'S30F, C48V, E172K', '+1.41', '探索性扫描'],
        ['2', 'ppluGFP', '4.300', 'E32V, T38E, L125T, M218L, V219Y, V220S, L221H', '-0.517', '分层突变'],
        ['3', 'amacGFP', '3.985', 'I14V, Q157A, K166T, T203R, T230Q', '+2.039', '分层突变'],
        ['4', 'amacGFP', '3.982', 'V11L, R122V, K162A', '-1.840', '探索性扫描'],
        ['5', 'avGFP', '3.709', 'M78R, G116N, F165E', '-0.878', '探索性扫描'],
    ]
)

# ======================== 10. 终极汇总 ========================
add_heading('10. 终极汇总管线 (final_pipeline)', level=1)
add_para('对应程序: final_pipeline.py', indent=False)
add_para('final_pipeline.py是项目的最终汇总脚本,将多个数据源的候选序列合并后统一打分:', indent=True)
add_para('(1) 合并所有Phase 2产出(含多轮迭代结果)及交叉验证数据。')
add_para('(2) 对缺失分数的序列补充V5/V7/V9三模型打分。')
add_para('(3) 重算共识投票,按类型独立阈值筛选。')
add_para('(4) 每种GFP类型选最佳1条 + 第5条取跨类型最高分,产出最终5条候选。')

# ======================== 11. 模型与工具来源 ========================
add_heading('11. 模型与工具来源', level=1)

add_table(
    ['模型/工具', '来源', '地址'],
    [
        ['ESM-2 t33 650M', 'Meta AI (Facebook)', 'GitHub: github.com/facebookresearch/esm\nHuggingFace: huggingface.co/facebook/esm2_t33_650M_UR50D'],
        ['ESM3 (开源版)', 'Evolutionary Scale', 'GitHub: github.com/evolutionaryscale/esm'],
        ['ESM3-ddG v2', 'Hazem Essam', 'HuggingFace: huggingface.co/hazemessam/esm3_ddg_v2'],
        ['PEFT (LoRA)', 'Hugging Face', 'GitHub: github.com/huggingface/peft'],
        ['AlphaFold 3', 'Google DeepMind', 'GitHub: github.com/google-deepmind/alphafold3'],
        ['Transformers', 'Hugging Face', 'GitHub: github.com/huggingface/transformers'],
        ['PyTorch', 'Meta AI', 'pytorch.org'],
    ]
)

# ======================== 12. 辅助工具 ========================
add_heading('12. 辅助工具与验证程序', level=1)

add_heading('12.1 散点图工具', level=2)
add_para('scatter_v7.py / scatter_v9.py: 生成模型预测散点图,按亮度区间均匀采样,展示各GFP类型的R2和MAE。历史散点图(V4/V5)存放在prediction comparison/目录。', indent=True)

add_heading('12.2 湿实验序列验证', level=2)
add_para('test_wetlab_seqs.py: 对20条已知高亮度序列用三模型打分,人工验证模型预测是否合理。不能盲目信任R2,需确保模型在真实序列上给出合理预测。', indent=True)

add_heading('12.3 单序列预测', level=2)
add_para('test_model.py / predict_v9.py: 交互式输入氨基酸序列,获得三模型预测亮度值,快速验证新设计序列。', indent=True)

# ======================== 13. 项目文件清单 ========================
add_heading('13. 项目文件清单', level=1)

add_heading('13.1 核心代码 (src/)', level=2)
files_src = [
    ('config.py', '全局配置:路径、常量、随机种子'),
    ('phase1_data.py', 'Phase 1: 数据清洗与ESM-2特征提取'),
    ('phase1_lora_finetune_v5.py', 'Phase 1: V5模型训练(全量LoRA + 122k)'),
    ('phase1_lora_finetune_v6.py', 'Phase 1: V9模型训练(LoRA + 72k + 深层MLP头)'),
    ('phase1_frozen_esm2_mlp.py', 'Phase 1: V7模型训练(冻结ESM-2 + 残差MLP)'),
    ('phase2_generate_hierarchical.py', 'Phase 2: 分层突变 + 共识筛选(5轮迭代)'),
    ('phase3_ddg_scoring.py', 'Phase 3: ESM3-ddG热稳定性打分'),
    ('phase4_scoring.py', 'Phase 4: 综合筛选 -> Top 20'),
    ('prepare_72k_data.py', '72k优质数据分层筛选'),
]
for name, desc in files_src:
    add_para(f'  {name} — {desc}')

add_heading('13.2 根目录核心脚本', level=2)
files_root = [
    ('main.py', '统一命令行入口'),
    ('final_pipeline.py', '终极管线:四数据源合并+三模型打分+选5条'),
    ('test_model.py', '单序列亮度预测(交互式)'),
    ('predict_v9.py', 'V9单序列预测'),
    ('scatter_v7.py / scatter_v9.py', '模型散点图'),
    ('test_wetlab_seqs.py', '湿实验序列验证'),
]
for name, desc in files_root:
    add_para(f'  {name} — {desc}')

add_heading('13.3 数据文件', level=2)
add_para('  data/raw/GFP_data.xlsx — 原始实验数据(141,572条 + beforetopseqs 20条)')
add_para('  data/raw/Exclusion_List.csv — 排除列表(135,414条)')
add_para('  generated_seqs/all_scored_merged_HCX_260625.csv — 39,020条全量打分')
add_para('  generated_seqs/final_5_candidates_v3_HCX_260625.csv — 最终5条候选')
add_para('  generated_seqs/phase2_full_scored_HCX_260623.csv — 分层突变完整结果(6,626条)')
add_para('  generated_seqs/iterations/ — 各轮迭代详细数据(iter1~iter5)')
add_para('  generated_seqs/intermediate/ — 中间筛选结果(V7过滤、交叉验证、ddG打分、Top 20)')
add_para('  generated_seqs/archive/ — 旧方案数据及历史版本')

add_heading('13.4 模型存档', level=2)
add_para('  ESM2_LoRA_V9_72k/ — V9主力(R²=0.946)')
add_para('  ESM2_Frozen_MLP_V7_72k/ — V7初筛(R²=0.921)')
add_para('  ESM2_LoRA_V5_TypeEmbed/ — V5安全网')

add_heading('13.5 归档文件', level=2)
add_para('历代迭代代码、调试脚本、一次性分析工具已打包至 archive/ 目录。包含V1-V4传统模型训练代码、各轮迭代的分析脚本(check_*、diagnose_*等)以及候选替换中间版本。', indent=True)

# ======================== 14. 参考文献 ========================
add_heading('14. 参考文献', level=1)
refs = [
    '[1] Lin et al. Evolutionary-scale prediction of atomic-level protein structure with a large language model. Science, 2023.',
    '[2] Hayes et al. Simulating 500 million years of evolution with a language model. bioRxiv, 2024.',
    '[3] Alley et al. Unified rational protein engineering with sequence-based deep representation learning. Nature Methods, 2019.',
    '[4] Hu et al. LoRA: Low-Rank Adaptation of Large Language Models. ICLR, 2022.',
    '[5] Jumper et al. Highly accurate protein structure prediction with AlphaFold. Nature, 2021.',
]
for ref in refs:
    add_para(ref)

# ======================== 保存 ========================
out_path = 'd:/GFP_Hexagon_Warrior/docs/GFP_Hexagon_Warrior_技术文档_v9.docx'
doc.save(out_path)
print(f'文档已保存: {out_path}')
